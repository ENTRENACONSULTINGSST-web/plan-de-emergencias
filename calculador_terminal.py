#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Calculador de Vulnerabilidad por Rombos (Edición de Terminal - Sin requerir Node/NPM)
Permite realizar la evaluación completa de SG-SST según la metodología del Diamante de Riesgos.
Genera reportes técnicos oficiales en Word (.docx), Excel (.xlsx) e imágenes de alta definición.
Complementario con repositorio GitHub: https://github.com
"""

import os
import sys
import json
import math
from datetime import datetime

# Estilos de formato ANSI para el terminal
class Color:
    GREEN = '\033[92m'
    YELLOW = '\033[93m'
    RED = '\033[91m'
    CYAN = '\033[96m'
    MAGENTA = '\033[95m'
    WHITE = '\033[97m'
    GRAY = '\033[90m'
    BOLD = '\033[1m'
    UNDERLINE = '\033[4m'
    RESET = '\033[0m'

# Datos Maestros de la Metodología
COMPONENT_BLOCKS = [
    {
        'id': 'Personas',
        'name': 'Vulnerabilidad en Personas',
        'items': [
            {
                'id': 'P_Organizacion',
                'name': 'Organización',
                'description': 'Estructura organizativa, comités de emergencia, brigadas y asignación presupuestal.',
                'questions': [
                    {
                        'id': 'p_org_q1',
                        'text': '¿Existe un Comité Operativo de Emergencias (COE) conformado y con roles asignados por escrito?',
                        'recommendation': 'Conformar el COE mediante acta de gerencia, definiendo líder, enlace y responsabilidades ante emergencias.'
                    },
                    {
                        'id': 'p_org_q2',
                        'text': '¿Se cuenta con una Brigada de Emergencia debidamente seleccionada, entrenada y estructurada?',
                        'recommendation': 'Elegir personal apto físicamente y capacitarlo legalmente. Nombrar jefe de brigada y coordinadores de evacuación.'
                    },
                    {
                        'id': 'p_org_q3',
                        'text': '¿Se cuenta con un presupuesto anual asignado y aprobado específicamente para el plan de prevención y emergencias?',
                        'recommendation': 'Asignar rubro presupuestal anual certificado para la compra de suministros, recargas de extintores y capacitación.'
                    }
                ]
            },
            {
                'id': 'P_Capacitacion',
                'name': 'Capacitación',
                'description': 'Capacitación teórico-práctica del personal, formación a la brigada y divulgación general.',
                'questions': [
                    {
                        'id': 'p_cap_q1',
                        'text': '¿Existe un plan formativo teórico-práctico anual para el personal general, contratistas y brigadistas?',
                        'recommendation': 'Elaborar un cronograma anual que abarque inducción de emergencias a nuevos ingresos y reentrenamiento a antiguos.'
                    },
                    {
                        'id': 'p_cap_q2',
                        'text': '¿Se cuenta con entrenamiento técnico específico para la brigada en rescate, primeros auxilios y control de incendios?',
                        'recommendation': 'Garantizar prácticas presenciales certificadas con bomberos u organismos de socorro especializados.'
                    },
                    {
                        'id': 'p_cap_q3',
                        'text': '¿Se realizan campaigns de divulgación periódica del plan de emergencias, mapas y puntos de encuentro a los colaboradores?',
                        'recommendation': 'Utilizar correos informativos, carteleras, folletos y charlas de 5 minutos sobre rutas y puntos de encuentro.'
                    }
                ]
            },
            {
                'id': 'P_Dotacion',
                'name': 'Dotación',
                'description': 'Equipos de Protección Personal, realización de simulacros organizados y sistemas de comunicación.',
                'questions': [
                    {
                        'id': 'p_dot_q1',
                        'text': '¿Los integrantes de la brigada cuentan con Equipos de Protección Personal (EPP) completos, normados e idóneos?',
                        'recommendation': 'Suministrar cascos, guantes, goggles, chalecos de alta visibilidad y botas de seguridad a toda la brigada.'
                    },
                    {
                        'id': 'p_dot_q2',
                        'text': '¿Se realizan simulacros de evacuación planificados (mínimo 1 anual por ley) con evaluación formal de tiempos de respuesta?',
                        'recommendation': 'Programar, ejecutar y redactar el informe técnico del simulacro nacional anual, midiendo tiempos con cronómetro.'
                    },
                    {
                        'id': 'p_dot_q3',
                        'text': '¿La organización dispone de sistemas de comunicación de emergencias idóneos (radios bidireccionales, megáfonos, alarmas)?',
                        'recommendation': 'Adquirir al menos un megáfono portátil, radios con frecuencia reservada y un sistema sonoro audible en toda la instalación.'
                    }
                ]
            }
        ]
    },
    {
        'id': 'Recursos',
        'name': 'Vulnerabilidad en Recursos',
        'items': [
            {
                'id': 'R_Materiales',
                'name': 'Materiales',
                'description': 'Clasificación de botiquines, camillas de respuesta, y kits especiales para control de eventos.',
                'questions': [
                    {
                        'id': 'r_mat_q1',
                        'text': '¿Se cuenta con botiquines de primeros auxilios debidamente dotados, vigentes e inspeccionados según la norma local?',
                        'recommendation': 'Sellar los botiquines, inventariar mensualmente los insumos y descartar elementos vencidos.'
                    },
                    {
                        'id': 'r_mat_q2',
                        'text': '¿Se dispone de camillas de emergencia con inmovilizadores y correas de sujeción (araña) distribuidas estratégicamente?',
                        'recommendation': 'Instalar camillas rígidas en la pared, con sus respectivos inmovilizadores de cuello y correas correspondientes.'
                    },
                    {
                        'id': 'r_mat_q3',
                        'text': '¿Se cuenta con kits para derrames químicos, biológicos o tecnológicos completos (absorbentes, diques, bolsas rojas)?',
                        'recommendation': 'Adquirir kits de derrames en puntos de almacenamiento de químicos con absorbentes y pala plástica.'
                    }
                ]
            },
            {
                'id': 'R_Infraestructura',
                'name': 'Infraestructura',
                'description': 'Ubicación y estado de rutas de evacuación, sismorresistencia del predio y diseño seguro de escaleras.',
                'questions': [
                    {
                        'id': 'r_inf_q1',
                        'text': '¿Las rutas de evacuación, salidas de emergencia y pasillos están despejados, señalizados e iluminados autónomamente?',
                        'recommendation': 'Garantizar que no existan mercancías obstruyendo pasillos. Instalar lámparas de emergencia certificadas.'
                    },
                    {
                        'id': 'r_inf_q2',
                        'text': '¿La estructura de la edificación cuenta con concepto técnico de sismorresistencia (NSR-10) o reforzamiento estructural?',
                        'recommendation': 'Contratar una firma de ingeniería civil para realizar estudio técnico analítico y emitir concepto legal de sismicidad estructural.'
                    },
                    {
                        'id': 'r_inf_q3',
                        'text': '¿Las escaleras de contra incendios cuentan con pasamanos continuos, cintas antideslizantes de alta tracción y material ignífugo?',
                        'recommendation': 'Fijar pasamanos laterales robustos y reemplazar cintas de lija desgastadas en los cantos de las gradas.'
                    }
                ]
            },
            {
                'id': 'R_Equipos',
                'name': 'Equipos',
                'description': 'Capacidad extintora portátil, detección automatizada de incendios y sistemas fijos de extinción.',
                'questions': [
                    {
                        'id': 'r_equ_q1',
                        'text': '¿Se cuenta con extintores portátiles multipropósito vigentes, anclados a altura reglamentaria, señalizados e inspeccionados?',
                        'recommendation': 'Realizar inspección visual de presión mensual. Programar recarga anual obligatoria e instalarlos a una altura de 1.2 a 1.5 metros.'
                    },
                    {
                        'id': 'r_equ_q2',
                        'text': '¿La instalación cuenta con un sistema electrónico de detección de humo, calor o gas completamente funcional y centralizado?',
                        'recommendation': 'Instalar panel de alarma centralizado conectado a detectores fotoeléctricos con mantenimiento preventivo semestral.'
                    },
                    {
                        'id': 'r_equ_q3',
                        'text': '¿Existen sistemas fijos de extinción de incendios (gabinetes equipados con manguera, rociadores automáticos) certificados?',
                        'recommendation': 'Verificar la presión del sistema hidroneumático de la red contra incendio y capacitar al personal pesado en su despliegue.'
                    }
                ]
            }
        ]
    },
    {
        'id': 'Sistemas',
        'name': 'Vulnerabilidad en Sistemas y Procesos',
        'items': [
            {
                'id': 'S_Servicios',
                'name': 'Servicios Públicos',
                'description': 'Integridad de los servicios de energía, redes hidrosanitarias y seguridad del gas natural/propano.',
                'questions': [
                    {
                        'id': 's_ser_q1',
                        'text': '¿El suministro de energía eléctrica posee cableado en buen estado, cajas tabuladas rotuladas e inspecciones termográficas?',
                        'recommendation': 'Inspeccionar cableado expuesto periódicamente. Mantener cerrados los armarios eléctricos bajo norma RETIE.'
                    },
                    {
                        'id': 's_ser_q2',
                        'text': '¿Se cuenta con servicios independientes de agua potable y alcantarillado con mantenimiento sistemático preventivo?',
                        'recommendation': 'Limpiar trampas de grasas e inspeccionar cajas de paso sanitario de forma semestral para prevenir taponamientos.'
                    },
                    {
                        'id': 's_ser_q3',
                        'text': '¿Las redes de gas natural o propano se encuentran certificadas bajo norma, con tuberías pintadas de amarillo y válvula exterior?',
                        'recommendation': 'Instalar válvula de paso rápido en el exterior con señalización de emergencia para corte rápido.'
                    }
                ]
            },
            {
                'id': 'S_Alternos',
                'name': 'Sistemas Alternos',
                'description': 'Suministro alterno de agua, generación de potencia de respaldo y resguardo de datos.',
                'questions': [
                    {
                        'id': 's_alt_q1',
                        'text': '¿Dispone la empresa de tanques de reserva de agua con capacidad autónoma calculada para abastecimiento ante cortes prolongados?',
                        'recommendation': 'Calcular el consumo de agua por colaborador y realizar lavado/desinfección semestral obligatoria de tanques.'
                    },
                    {
                        'id': 's_alt_q2',
                        'text': '¿Se cuenta con una planta eléctrica de emergencia operativa, con transferencia automática y reserva de combustible segura?',
                        'recommendation': 'Garantizar arranque de prueba semanal de la planta bajo carga y almacenar diésel en cubas estancas.'
                    },
                    {
                        'id': 's_alt_q3',
                        'text': '¿Se tienen respaldos (backups) diarios automatizados de información crítica del negocio almacenados fuera de la sede?',
                        'recommendation': 'Configurar respaldos en nube cifrada sincronizados diariamente de manera remota automatizada.'
                    }
                ]
            },
            {
                'id': 'S_Recuperacion',
                'name': 'Sistemas de Recuperación',
                'description': 'Suscripción de pólizas ante catástrofe, convenios críticos de soporte y continuidad del negocio.',
                'questions': [
                    {
                        'id': 's_rec_q1',
                        'text': '¿Se encuentra suscrita y plenamente vigente la Póliza de Seguros Todo Riesgo con cobertura explícita de desastres naturales?',
                        'recommendation': 'Validar cobertura anual frente a sismos, asonadas e incendios, ajustando topes según el valor real de activos.'
                    },
                    {
                        'id': 's_rec_q2',
                        'text': '¿Se tienen convenios o contratos de emergencia vigentes con proveedores claves de agua, alimentos, maquinaria y salud?',
                        'recommendation': 'Definir acuerdos firmados y canales con centros de salud locales y distribuidores de material de contingencia.'
                    },
                    {
                        'id': 's_rec_q3',
                        'text': '¿Se cuenta con un Plan de Continuidad del Negocio (BCP) documentado, socializado y aprobado formalmente por la gerencia?',
                        'recommendation': 'Crear el manual de continuidad determinando tiempos de restablecimiento operativo para procesos críticos.'
                    }
                ]
            }
        ]
    }
]

TRANSVERSAL_BLOCKS = [
    {
        'id': 'PESV',
        'name': 'Plan Estratégico de Seguridad Vial (PESV)',
        'lawReference': 'Resolución 20223040040595 de 2022 (Colombia)',
        'description': 'Gestión de riesgos en la vía para empresas con flota de vehículos o conductores misionales.',
        'questions': [
            {
                'id': 'pesv_q1',
                'text': '¿Se cuenta con un diagnóstico detallado de riesgos viales y caracterización sociodemográfica de todos los conductores?',
                'recommendation': 'Aplicar censo y encuesta vial para catalogar perfiles de conducción, estado de licencias y horarios.'
            },
            {
                'id': 'pesv_q2',
                'text': '¿Se realizan y registran de forma obligatoria las inspecciones preoperacionales a todos los vehículos antes de cada jornada?',
                'recommendation': 'Digitalizar planillas diarias evaluando el estado físico de frenos, fluidos, llantas y dirección.'
            },
            {
                'id': 'pesv_q3',
                'text': '¿El personal motorizado o conductor posee certificaciones vigentes en cursos de manejo defensivo y seguridad vial?',
                'recommendation': 'Programar reentrenamientos específicos de teoría y práctica de conducción preventiva con entes aprobados.'
            },
            {
                'id': 'pesv_q4',
                'text': '¿El PESV se encuentra implementado bajo la estructura normativa vigente y articulado con el SGSST de la empresa?',
                'recommendation': 'Vincular el Comité de Seguridad Vial con el COPASST y auditar anualmente las líneas de acción.'
            }
        ]
    },
    {
        'id': 'Psicosocial',
        'name': 'Riesgo Psicosocial y Salud Mental',
        'lawReference': 'Resolución 2764 de 2022 y Ley 1010 de 2006 (Colombia)',
        'description': 'Control de fatiga, prevención del estrés e idoneidad del Comité de Convivencia Laboral.',
        'questions': [
            {
                'id': 'psico_q1',
                'text': '¿Se aplica la Batería oficial de Riesgo Psicosocial de forma anual por un psicólogo titulado con licencia en SST?',
                'recommendation': 'Contratar psicólogo certificado para administrar la prueba ministerial y redactar el reporte de intervención obligatorio.'
            },
            {
                'id': 'psico_q2',
                'text': '¿El Comité de Convivencia Laboral (COCOLA) se encuentra constituido de forma de paridad, activo y con actas trimestrales?',
                'recommendation': 'Garantizar elecciones de representantes para el COCOLA, manteniendo las debidas actas de resolución mensuales/trimestrales.'
            },
            {
                'id': 'psico_q3',
                'text': '¿Dispone la empresa de una ruta de atención, protocolo o línea de primeros auxilios psicológicos ante crisis?',
                'recommendation': 'Elaborar un flujograma institucional y divulgar las líneas de orientación telefónica pública de soporte psiquiátrico.'
            },
            {
                'id': 'psico_q4',
                'text': '¿Se implementan actividades continuas de prevención de Burnout, medición de clima y balance de carga de trabajo?',
                'recommendation': 'Estructurar políticas para respeto de la desconexión digital e incentivar pausas activas cognitivas diarias.'
            }
        ]
    },
    {
        'id': 'Alturas',
        'name': 'Trabajo Seguro en Alturas',
        'lawReference': 'Resolución 4272 de 2021 (Colombia)',
        'description': 'Gestión del riesgo en trabajos realizados por encima de 2.0 metros de altura.',
        'questions': [
            {
                'id': 'alt_q1',
                'text': '¿Se cuenta con el programa de prevención y protección contra caídas firmado por un Coordinador de Alturas?',
                'recommendation': 'Elaborar y actualizar anualmente el programa estipulando el listado completo de áreas, anclajes permanentes y controles.'
            },
            {
                'id': 'alt_q2',
                'text': '¿Los trabajadores que ejecutan labores en alturas cuentan con competencias vigentes y aptitud de examen médico?',
                'recommendation': 'Hacer seguimiento estricto de las licencias formativas vigentes y agendar exámenes de control clínico ocupacional (perfil lipídico, vértigo).'
            },
            {
                'id': 'alt_q3',
                'text': '¿Los puntos de anclaje, líneas de vida y arneses cuentan con certificación física de fabricante y prueba anual?',
                'recommendation': 'Disponer de expedientes individuales para cada arnés e inspeccionarlos físicamente antes de cada labor de izaje.'
            },
            {
                'id': 'alt_q4',
                'text': '¿Se tiene estipulado un procedimiento escrito para rescate en alturas y se cuenta con al menos un kit de rescate?',
                'recommendation': 'Adquirir un kit de descenso rápido dotado y capacitar a la brigada en protocolos de rescate auto-retractil.'
            }
        ]
    },
    {
        'id': 'Litio',
        'name': 'Riesgo Tecnológico en Baterías de Litio',
        'lawReference': 'Buenas Prácticas NFPA 855 y Gestión Tecnológica SST (Colombia)',
        'description': 'Prevención de embalamiento térmico en zonas de carga de montacargas, scooters o sistemas UPS.',
        'questions': [
            {
                'id': 'litio_q1',
                'text': '¿Las zonas de carga de equipos con baterías de litio se encuentran en áreas independientes, ventiladas y sin combustibles?',
                'recommendation': 'Establecer una jaula de carga retirada de papelería, maderas o bodegas principales, con extracción mecánica de aire activa.'
            },
            {
                'id': 'litio_q2',
                'text': '¿Se dispone de agentes extintores encapsuladores de agua con aditivo especial (ej: F-500) óptimos para mitigar fuegos químicos clase D/Litio?',
                'recommendation': 'Instalar extintores cargados con agente F-500 en las bahías de carga para detener la reacción exotérmica de embalamiento químico.'
            },
            {
                'id': 'litio_q3',
                'text': '¿Se realiza inspección térmica continua (termografía o pirómetros digitales rápidos) sobre los cargadores de litio en horas pico?',
                'recommendation': 'Implementar rondas con termómetros láser apuntando a las conexiones e incorporar sensores de sobretemperatura con corte automático.'
            },
            {
                'id': 'litio_q4',
                'text': '¿Se encuentra documentado un protocolo estricto para aislamiento o sumersión rápida de celdas infladas, húmedas o calientes?',
                'recommendation': 'Tener dispuesto un contenedor lleno de arena seca o agua desmineralizada para sumergir de inmediato una unidad sospechosa antes del escape térmico.'
            }
        ]
    },
    {
        'id': 'Quimicos',
        'name': 'Riesgo Químico - Sistema Globalmente Armonizado',
        'lawReference': 'Decreto 1496 de 2018 y Resoluciones SGA (Colombia)',
        'description': 'Uso de pictogramas, fichas de seguridad en puestos de trabajo, matrices de compatibilidad física y diques.',
        'questions': [
            {
                'id': 'quim_q1',
                'text': '¿Todos los productos químicos de la edificación están marcados bajo el Sistema Globalmente Armonizado (SGA) con pictogramas claros?',
                'recommendation': 'Etiquetar atomizadores de limpieza, lubricantes o solventes con las palabras de advertencia (Peligro/Atención) e indicaciones de daño.'
            },
            {
                'id': 'quim_q2',
                'text': '¿Las Fichas de Datos de Seguridad (FDS) con estructura de 16 secciones se encuentran impresas o cargadas para consulta en el sitio?',
                'recommendation': 'Imprimir las FDS de los productos de mayor uso y colocarlas en carpetas de acrílico junto a las estaciones donde se manipulan.'
            },
            {
                'id': 'quim_q3',
                'text': '¿Existe y se respeta físicamente una matriz de compatibilidad de almacenamiento de sustancias químicas en la bodega?',
                'recommendation': 'Separar físicamente ácidos de bases, y combustibles de oxidantes mediante distancias mínimas de 3 metros o barreras físicas.'
            },
            {
                'id': 'quim_q4',
                'text': '¿Los tanques o canecas de gran volumen cuentan con diques o bandejas de contención con capacidad del 110% del envase principal?',
                'recommendation': 'Construir diques de mampostería o adquirir estibas recolectoras plásticas de derrames certificadas para contener el volumen total.'
            }
        ]
    }
]

STANDARD_THREATS = [
    {
        'id': 'incendio',
        'name': 'Incendio Estructural / Explosión',
        'level': 'POSIBLE',
        'description': 'Ignición súbita por fallas eléctricas en subestaciones, fugas de combustible o almacenamiento inadecuado de sólidos inflamables.'
    },
    {
        'id': 'sismo',
        'name': 'Sismo / Terremoto',
        'level': 'PROBABLE',
        'description': 'Movimiento telúrico severo derivado de fallas geológicas activas (Colombia se encuentra en zona de alta sismicidad).'
    },
    {
        'id': 'inundacion',
        'name': 'Inundación por Lluvias / Escorrentía',
        'level': 'POSIBLE',
        'description': 'Saturación de desagües fluviales públicos o desbordamiento de vertientes hídricas cercanas ante tormentas atípicas.'
    },
    {
        'id': 'derrame',
        'name': 'Derrame de Materiales Peligrosos',
        'level': 'POSIBLE',
        'description': 'Pérdida de contención en tanques de ACPM, tanques de amoníaco de refrigeración o manipulación deficiente de ácidos.'
    },
    {
        'id': 'riesgo_publico',
        'name': 'Hurto, Asonada o Terrorismo (Riesgo Público)',
        'level': 'POSIBLE',
        'description': 'Alteraciones de orden público en inmediaciones del predio corporativo, vandalismo o ingreso no autorizado.'
    },
    {
        'id': 'litio_explosion',
        'name': 'Embalamiento Térmico de Baterías de Litio',
        'level': 'POSIBLE',
        'description': 'Fuego químico incontrolable originado en el rack de recarga de montacargas, patinetas o dispositivos de flota.'
    }
]

EVALUATION_PRESETS = [
    {
        'name': 'Empresa Industrial Certificada (Cumplimiento Alto 🟢)',
        'threat_id': 'sismo',
        'threat_name': 'Sismo / Terremoto',
        'threat_desc': 'Movimiento telúrico severo derivado de fallas geológicas activas.',
        'threat_level': 'PROBABLE',
        'answers': {
            'p_org_q1': 1.0, 'p_org_q2': 1.0, 'p_org_q3': 1.0,
            'p_cap_q1': 1.0, 'p_cap_q2': 1.0, 'p_cap_q3': 1.0,
            'p_dot_q1': 1.0, 'p_dot_q2': 0.5, 'p_dot_q3': 1.0,
            'r_mat_q1': 1.0, 'r_mat_q2': 1.0, 'r_mat_q3': 1.0,
            'r_inf_q1': 1.0, 'r_inf_q2': 0.5, 'r_inf_q3': 1.0,
            'r_equ_q1': 1.0, 'r_equ_q2': 1.0, 'r_equ_q3': 1.0,
            's_ser_q1': 1.0, 's_ser_q2': 1.0, 's_ser_q3': 1.0,
            's_alt_q1': 1.0, 's_alt_q2': 0.5, 's_alt_q3': 1.0,
            's_rec_q1': 1.0, 's_rec_q2': 1.0, 's_rec_q3': 1.0,
            'pesv_q1': 1.0, 'pesv_q2': 1.0, 'pesv_q3': 1.0, 'pesv_q4': 1.0,
            'psico_q1': 1.0, 'psico_q2': 1.0, 'psico_q3': 0.5, 'psico_q4': 0.5,
            'alt_q1': 1.0, 'alt_q2': 1.0, 'alt_q3': 1.0, 'alt_q4': 0.5,
            'litio_q1': 1.0, 'litio_q2': 1.0, 'litio_q3': 0.5, 'litio_q4': 1.0,
            'quim_q1': 1.0, 'quim_q2': 1.0, 'quim_q3': 1.0, 'quim_q4': 0.5,
        }
    },
    {
        'name': 'Establecimiento Comercial PyME (Vulnerabilidad Media 🟡)',
        'threat_id': 'incendio',
        'threat_name': 'Incendio Estructural / Explosión',
        'threat_desc': 'Ignición por fallas eléctricas en subestaciones o almacenamiento inadecuado.',
        'threat_level': 'PROBABLE',
        'answers': {
            'p_org_q1': 0.5, 'p_org_q2': 0.5, 'p_org_q3': 0.5,
            'p_cap_q1': 0.5, 'p_cap_q2': 0.5, 'p_cap_q3': 1.0,
            'p_dot_q1': 0.5, 'p_dot_q2': 0.5, 'p_dot_q3': 0.0,
            'r_mat_q1': 1.0, 'r_mat_q2': 0.5, 'r_mat_q3': 0.0,
            'r_inf_q1': 0.5, 'r_inf_q2': 0.0, 'r_inf_q3': 0.5,
            'r_equ_q1': 1.0, 'r_equ_q2': 0.5, 'r_equ_q3': 0.0,
            's_ser_q1': 0.5, 's_ser_q2': 1.0, 's_ser_q3': 0.5,
            's_alt_q1': 0.5, 's_alt_q2': 0.0, 's_alt_q3': 0.5,
            's_rec_q1': 0.5, 's_rec_q2': 0.5, 's_rec_q3': 0.0,
            'pesv_q1': 0.5, 'pesv_q2': 0.5, 'pesv_q3': 0.0, 'pesv_q4': 0.0,
            'psico_q1': 0.5, 'psico_q2': 1.0, 'psico_q3': 0.0, 'psico_q4': 0.5,
            'alt_q1': 0.5, 'alt_q2': 0.5, 'alt_q3': 0.0, 'alt_q4': 0.0,
            'litio_q1': 0.5, 'litio_q2': 0.0, 'litio_q3': 0.0, 'litio_q4': 0.0,
            'quim_q1': 1.0, 'quim_q2': 0.5, 'quim_q3': 0.5, 'quim_q4': 0.0,
        }
    },
    {
        'name': 'Edificio Crítico / Planta Tradicional (Vulnerabilidad Alta 🔴)',
        'threat_id': 'litio_explosion',
        'threat_name': 'Embalamiento Térmico de Baterías de Litio',
        'threat_desc': 'Fuego químico incontrolable originado en el rack de recarga de montacargas.',
        'threat_level': 'INMINENTE',
        'answers': {
            'p_org_q1': 0.0, 'p_org_q2': 0.0, 'p_org_q3': 0.0,
            'p_cap_q1': 0.0, 'p_cap_q2': 0.0, 'p_cap_q3': 0.5,
            'p_dot_q1': 0.0, 'p_dot_q2': 0.0, 'p_dot_q3': 0.0,
            'r_mat_q1': 0.5, 'r_mat_q2': 0.0, 'r_mat_q3': 0.0,
            'r_inf_q1': 0.0, 'r_inf_q2': 0.0, 'r_inf_q3': 0.0,
            'r_equ_q1': 0.5, 'r_equ_q2': 0.0, 'r_equ_q3': 0.0,
            's_ser_q1': 0.0, 's_ser_q2': 0.5, 's_ser_q3': 0.0,
            's_alt_q1': 0.0, 's_alt_q2': 0.0, 's_alt_q3': 0.0,
            's_rec_q1': 0.0, 's_rec_q2': 0.0, 's_rec_q3': 0.0,
            'pesv_q1': 0.0, 'pesv_q2': 0.0, 'pesv_q3': 0.0, 'pesv_q4': 0.0,
            'psico_q1': 0.0, 'psico_q2': 0.0, 'psico_q3': 0.0, 'psico_q4': 0.0,
            'alt_q1': 0.0, 'alt_q2': 0.0, 'alt_q3': 0.0, 'alt_q4': 0.0,
            'litio_q1': 0.0, 'litio_q2': 0.0, 'litio_q3': 0.0, 'litio_q4': 0.0,
            'quim_q1': 0.5, 'quim_q2': 0.0, 'quim_q3': 0.0, 'quim_q4': 0.0,
        }
    }
]

SESSION_FILE = "respuestas_sesion.json"

class EvaluacionEngine:
    def __init__(self):
        self.answers = {}
        self.threat_name = "Amenaza Desconocida"
        self.threat_desc = "Falta caracterización del escenario."
        self.threat_level = "POSIBLE" # POSIBLE, PROBABLE, INMINENTE

    def load_preset(self, index):
        p = EVALUATION_PRESETS[index]
        self.threat_name = p['threat_name']
        self.threat_desc = p['threat_desc']
        self.threat_level = p['threat_level']
        self.answers = p['answers'].copy()

    def set_answer(self, q_id, val):
        self.answers[q_id] = val

    def get_vulnerability_details(self, avg):
        if avg <= 0.40:
            return {'level': 'ALTA', 'color': 'red', 'color_ansi': Color.RED}
        elif avg <= 0.60:
            return {'level': 'MEDIA', 'color': 'yellow', 'color_ansi': Color.YELLOW}
        else:
            return {'level': 'BAJA', 'color': 'green', 'color_ansi': Color.GREEN}

    def calc_results(self):
        # 1. Calcular Core
        core_results = {}
        for block in COMPONENT_BLOCKS:
            block_sum = 0
            item_averages = {}
            for item in block['items']:
                item_sum = 0
                for q in item['questions']:
                    item_sum += self.answers.get(q['id'], 0.0) # default missing is 0 (Malo)
                item_avg = item_sum / len(item['questions']) if len(item['questions']) > 0 else 0
                item_averages[item['id']] = round(item_avg, 2)
                block_sum += item_avg
            block_avg = block_sum / len(block['items']) if len(block['items']) > 0 else 0
            block_avg = round(block_avg, 2)
            details = self.get_vulnerability_details(block_avg)
            core_results[block['id']] = {
                'name': block['name'],
                'average': block_avg,
                'level': details['level'],
                'color': details['color'],
                'color_ansi': details['color_ansi'],
                'item_averages': item_averages
            }

        # 2. Calcular Transversales
        transversal_results = {}
        for block in TRANSVERSAL_BLOCKS:
            sum_val = 0
            for q in block['questions']:
                sum_val += self.answers.get(q['id'], 0.0)
            avg = sum_val / len(block['questions']) if len(block['questions']) > 0 else 0
            avg = round(avg, 2)
            details = self.get_vulnerability_details(avg)
            transversal_results[block['id']] = {
                'name': block['name'],
                'lawReference': block['lawReference'],
                'average': avg,
                'level': details['level'],
                'color': details['color'],
                'color_ansi': details['color_ansi']
            }

        # 3. Contar rombos rojos de vulnerabilidad (Core)
        red_vulnerability_rombos = 0
        for b_id in ['Personas', 'Recursos', 'Sistemas']:
            if core_results[b_id]['color'] == 'red':
                red_vulnerability_rombos += 1

        # 4. Rombos en rojo totales (Vulnerabilidad + Amenaza)
        is_threat_red = self.threat_level == 'INMINENTE'
        total_red_rombos = red_vulnerability_rombos + (1 if is_threat_red else 0)

        # 5. Nivel consolidado de vulnerabilidad
        if total_red_rombos >= 3:
            consolidated_vul_level = 'ALTA'
            consolidated_color_ansi = Color.RED
        elif total_red_rombos >= 1:
            consolidated_vul_level = 'MEDIA'
            consolidated_color_ansi = Color.YELLOW
        else:
            consolidated_vul_level = 'BAJA'
            consolidated_color_ansi = Color.GREEN

        # 6. Intersección de riesgo global
        vul = consolidated_vul_level
        thr = self.threat_level
        risk_level = 'MEDIO'
        risk_color_ansi = Color.YELLOW
        risk_desc = "Existen fallas parciales. Diseñar acciones a mediano plazo."

        if vul == 'BAJA':
            if thr == 'POSIBLE':
                risk_level = 'BAJO'
                risk_color_ansi = Color.GREEN
                risk_desc = "Instalación altamente protegida. Estructura preventiva óptima."
            else:
                risk_level = 'MEDIO'
                risk_color_ansi = Color.YELLOW
                risk_desc = "Excelente preparación, pero la amenaza es alta o inminente. Mantener alerta."
        elif vul == 'MEDIA':
            if thr == 'INMINENTE':
                risk_level = 'ALTO'
                risk_color_ansi = Color.RED
                risk_desc = "Fallas organizativas unidas a un riesgo inminente de impacto físico inmediato."
            else:
                risk_level = 'MEDIO'
                risk_color_ansi = Color.YELLOW
                risk_desc = "Existen debilidades de preparación parcial. Requiere plan de mejora continuo."
        elif vul == 'ALTA':
            if thr == 'POSIBLE':
                risk_level = 'MEDIO'
                risk_color_ansi = Color.YELLOW
                risk_desc = "Vulnerabilidad severa en la sede. El peligro se mitiga por baja probabilidad de la amenaza."
            else:
                risk_level = 'ALTO'
                risk_color_ansi = Color.RED
                risk_desc = "Vulnerabilidad crítica alineada a una amenaza letal. Prioridad de intervención legal inmediata."

        return {
            'core_results': core_results,
            'transversal_results': transversal_results,
            'red_vulnerability_rombos': red_vulnerability_rombos,
            'total_red_rombos': total_red_rombos,
            'consolidated_vul_level': consolidated_vul_level,
            'consolidated_color_ansi': consolidated_color_ansi,
            'risk_level': risk_level,
            'risk_color_ansi': risk_color_ansi,
            'risk_desc': risk_desc
        }

    def save_session(self):
        data = {
            'threat_name': self.threat_name,
            'threat_desc': self.threat_desc,
            'threat_level': self.threat_level,
            'answers': self.answers
        }
        try:
            with open(SESSION_FILE, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=4, ensure_ascii=False)
            return True
        except Exception as e:
            print(f"Error al guardar sesión: {e}")
            return False

    def load_session(self):
        if not os.path.exists(SESSION_FILE):
            return False
        try:
            with open(SESSION_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
            self.threat_name = data.get('threat_name', 'Amenaza Desconocida')
            self.threat_desc = data.get('threat_desc', 'Sin descripción')
            self.threat_level = data.get('threat_level', 'POSIBLE')
            self.answers = data.get('answers', {})
            return True
        except Exception as e:
            print(f"Error al cargar sesión previa: {e}")
            return False

# Generador de Imágenes para el Diamante (Compatible con Pillow / PIL)
def generar_imagen_diamante_pil(engine, outfile="diamante_riesgo.png"):
    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError:
        print(f"\n[{Color.RED}AVISO{Color.RESET}] No se encuentra instalada la librería 'Pillow' (PIL). "
              f"Instálela ejecutando: {Color.BOLD}pip install pillow{Color.RESET}")
        return False

    results = engine.calc_results()
    
    # Colores HEX a RGB
    colors_rgb = {
        'POSIBLE': (16, 185, 129),   # Verde Emerald
        'PROBABLE': (245, 158, 11),   # Amarillo Amber
        'INMINENTE': (239, 68, 68),  # Rojo Rose
        'BAJA': (16, 185, 129),
        'MEDIA': (245, 158, 11),
        'ALTA': (239, 68, 68)
    }

    t_color = colors_rgb.get(engine.threat_level, (16, 185, 129))
    p_color = colors_rgb.get(results['core_results']['Personas']['level'], (16, 185, 129))
    r_color = colors_rgb.get(results['core_results']['Recursos']['level'], (16, 185, 129))
    s_color = colors_rgb.get(results['core_results']['Sistemas']['level'], (16, 185, 129))

    # Lienzo de 600x660 px con fondo pizarra (#1e293b / RGB: 30, 41, 59)
    img_width, img_height = 600, 660
    img = Image.new("RGB", (img_width, img_height), (30, 41, 59))
    draw = ImageDraw.Draw(img)

    # Coordenadas escaladas para el diamante (SVG original x2)
    # 1. TOP Rombo: Threat
    rombo_top = [(300, 30), (450, 180), (300, 330), (150, 180)]
    # 2. LEFT Rombo: Personas
    rombo_left = [(150, 180), (300, 330), (150, 480), (0, 330)]
    # 3. RIGHT Rombo: Recursos
    rombo_right = [(450, 180), (600, 330), (450, 480), (300, 330)]
    # 4. BOTTOM Rombo: Sistemas
    rombo_bottom = [(300, 330), (450, 480), (300, 630), (150, 480)]

    def draw_rombo(points, fill_color):
        draw.polygon(points, fill=fill_color)
        draw.polygon(points, outline=(15, 23, 42), width=5) # Línea divisoria robusta

    # Dibujar Polígonos
    draw_rombo(rombo_top, t_color)
    draw_rombo(rombo_left, p_color)
    draw_rombo(rombo_right, r_color)
    draw_rombo(rombo_bottom, s_color)

    # Dibujar Línea de Guías dadas
    draw.line((300, 0, 300, 660), fill=(71, 85, 105), width=2)
    draw.line((0, 330, 600, 330), fill=(71, 85, 105), width=2)

    # Textos sobre los rombos (Pillow fallback fuente por defecto si no hay ttf)
    # Cargamos fuentes truetype o usamos la simple por defecto
    font_large, font_small = None, None
    try:
        # Intenta usar fuentes TrueType del sistema si existen
        sys_fonts = ["arialbd.ttf", "calibrib.ttf", "DejaVuSans-Bold.ttf", "sans-serif"]
        for sf in sys_fonts:
            try:
                font_large = ImageFont.truetype(sf, 18)
                font_small = ImageFont.truetype(sf, 15)
                break
            except IOError:
                continue
    except Exception:
        pass

    def draw_centered_text(x, y, line1, line2, color=(255, 255, 255)):
        if font_large and font_small:
            # Medir y centrar
            w1 = draw.textlength(line1, font=font_large)
            draw.text((x - w1/2, y - 15), line1, fill=color, font=font_large)
            w2 = draw.textlength(line2, font=font_small)
            draw.text((x - w2/2, y + 10), line2, fill=color, font=font_small)
        else:
            # Fallback dibujo básico sin fuentes personalizadas
            draw.text((x - 30, y - 10), line1, fill=color)
            draw.text((x - 20, y + 10), line2, fill=color)

    # Dibujar textos de información en cuadrantes
    draw_centered_text(300, 130, "AMENAZA", engine.threat_level)
    draw_centered_text(150, 290, "PERSONAS", f"V: {engine.calc_results()['core_results']['Personas']['average']:.2f}")
    draw_centered_text(450, 290, "RECURSOS", f"V: {engine.calc_results()['core_results']['Recursos']['average']:.2f}")
    draw_centered_text(300, 450, "SISTEMAS", f"V: {engine.calc_results()['core_results']['Sistemas']['average']:.2f}")

    # Círculo central (Hub de control del Diamante)
    center_x, center_y = 300, 330
    radius = 35
    draw.ellipse((center_x - radius, center_y - radius, center_x + radius, center_y + radius), fill=(15, 23, 42), outline=(100, 116, 139), width=3)
    
    # Texto del indicador del centro
    red_rombos_v = results['red_vulnerability_rombos']
    text_c = f"V={red_rombos_v}R"
    if font_large:
        w_c = draw.textlength(text_c, font=font_large)
        draw.text((center_x - w_c/2, center_y - 10), text_c, fill=(56, 189, 248), font=font_large) # Sky-blue color
    else:
        draw.text((center_x - 15, center_y - 5), text_c, fill=(56, 189, 248))

    try:
        img.save(outfile, "PNG")
        return True
    except Exception as e:
        print(f"Error al escribir la imagen {outfile}: {e}")
        return False

# Exportador a Word usando python-docx
def exportar_word_docx(engine, outfile="Plan_Prevencion_Emergencias.docx"):
    try:
        import docx
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
    except ImportError:
        print(f"\n[{Color.RED}AVISO{Color.RESET}] No se encuentra instalada la librería 'python-docx'. "
              f"Instálela ejecutando: {Color.BOLD}pip install python-docx{Color.RESET}")
        return False

    results = engine.calc_results()
    doc = Document()

    # Configuración de márgenes
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)

    # Título Institucional
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("PLAN INSTITUCIONAL DE PREVENCIÓN Y EMERGENCIAS\n")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(15, 23, 42) # Slate 900

    run_subtitle = p_title.add_run("Metodología de Diamante / Conforme Decreto 1072 & Resolución 0312 / República de Colombia")
    run_subtitle.font.name = 'Arial'
    run_subtitle.font.size = Pt(10)
    run_subtitle.font.color.rgb = RGBColor(13, 148, 136) # Teal 600

    # Capítulo I
    h1 = doc.add_paragraph()
    r_h1 = h1.add_run("I. CARACTERIZACIÓN GENERAL Y AMENAZAS")
    r_h1.font.name = 'Arial'
    r_h1.font.size = Pt(14)
    r_h1.font.bold = True

    # Cuadro de datos iniciales
    table_init = doc.add_table(rows=1, cols=2)
    table_init.style = 'Light Shading Accent 1'
    hdr_cells = table_init.rows[0].cells
    hdr_cells[0].text = "Parámetro"
    hdr_cells[1].text = "Detalle y Caracterización"

    # Sombrear la cabecera
    shading_xml_header = parse_xml(r'<w:shd {} w:fill="0F172A"/>'.format(nsdecls('w')))
    hdr_cells[0]._tc.get_or_add_tcPr().append(shading_xml_header)
    shading_xml_header2 = parse_xml(r'<w:shd {} w:fill="0F172A"/>'.format(nsdecls('w')))
    hdr_cells[1]._tc.get_or_add_tcPr().append(shading_xml_header2)

    # Filas
    rows_data = [
        ("Nombre de la Sede / Sucursal", "Sede Corporativa Principal"),
        ("Fecha de Evaluación / Auditoría", str(datetime.now().strftime("%d/%m/%Y %H:%M"))),
        ("Amenaza Identificada Analizada", engine.threat_name),
        ("Nivel Cualitativo de Probabilidad", engine.threat_level),
        ("Detalle o Escenario Físico", engine.threat_desc)
    ]
    for param, val in rows_data:
        row = table_init.add_row()
        row.cells[0].text = param
        row.cells[1].text = val

    # Espacio
    doc.add_paragraph()

    # Capítulo II - Diamante de Riesgos
    h2 = doc.add_paragraph()
    r_h2 = h2.add_run("II. DIAGNÓSTICO DEL DIAMANTE DE RIESGO Y VULNERABILIDAD")
    r_h2.font.name = 'Arial'
    r_h2.font.size = Pt(14)
    r_h2.font.bold = True

    # Insertar Imagen si existe
    img_path = "diamante_riesgo.png"
    if not os.path.exists(img_path):
        generar_imagen_diamante_pil(engine, img_path)

    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = p_img.add_run()
        run_img.add_picture(img_path, width=Inches(3.5))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_cap = p_cap.add_run("Gráfico 1: Estructura Dinámica y Semáforo del Diamante de Riesgo.")
        r_cap.italic = True
        r_cap.font.size = Pt(9)
        r_cap.font.color.rgb = RGBColor(100, 116, 139)

    p_risk = doc.add_paragraph()
    r_risk_title = p_risk.add_run("\nRESULTADO INTERSECCIÓN COMPLETADO: ")
    r_risk_title.bold = True
    r_risk_val = p_risk.add_run(f"RIESGO {results['risk_level']}\n")
    r_risk_val.bold = True
    if results['risk_level'] == 'ALTO':
        r_risk_val.font.color.rgb = RGBColor(239, 68, 68)
    elif results['risk_level'] == 'MEDIO':
        r_risk_val.font.color.rgb = RGBColor(245, 158, 11)
    else:
        r_risk_val.font.color.rgb = RGBColor(16, 185, 129)

    r_risk_desc = p_risk.add_run(f"{results['risk_desc']}\n")
    r_risk_desc.italic = True

    # Resumen cuantitativo cores
    table_core = doc.add_table(rows=1, cols=4)
    table_core.style = 'Light Shading Accent 1'
    hdr_c = table_core.rows[0].cells
    hdr_c[0].text = "Componente de Vulnerabilidad"
    hdr_c[1].text = "Promedio Obtenido"
    hdr_c[2].text = "Calificación"
    hdr_c[3].text = "Resultado Semáforo"

    headers_colors = ["0F172A", "0F172A", "0F172A", "0F172A"]
    for i, col in enumerate(headers_colors):
        shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(col, nsdecls('w')))
        hdr_c[i]._tc.get_or_add_tcPr().append(shading)

    for b_id, b_val in results['core_results'].items():
        row = table_core.add_row()
        row.cells[0].text = b_val['name']
        row.cells[1].text = f"{b_val['average']:.2f}"
        row.cells[2].text = b_val['level']
        row.cells[3].text = "🔴 CRÍTICO" if b_val['color'] == 'red' else ("🟡 ADVERTENCIA" if b_val['color'] == 'yellow' else "🟢 PROTEGIDO")

    doc.add_paragraph()

    # Capítulo III - Módulos Transversales Especiales
    h3 = doc.add_paragraph()
    r_h3 = h3.add_run("III. PLANILLA DE REQUISITOS TRANSVERSALES ESPECIALES")
    r_h3.font.name = 'Arial'
    r_h3.font.size = Pt(14)
    r_h3.font.bold = True

    table_trans = doc.add_table(rows=1, cols=4)
    table_trans.style = 'Light Shading Accent 1'
    hdr_t = table_trans.rows[0].cells
    hdr_t[0].text = "Bloque Reglamentado"
    hdr_t[1].text = "Puntaje Promedio"
    hdr_t[2].text = "Vulnerabilidad"
    hdr_t[3].text = "Referencia Normativa de Ley"

    for i in range(4):
        shading = parse_xml(r'<w:shd {} w:fill="1F2937"/>'.format(nsdecls('w')))
        hdr_t[i]._tc.get_or_add_tcPr().append(shading)

    for b_id, b_val in results['transversal_results'].items():
        row = table_trans.add_row()
        row.cells[0].text = b_val['name']
        row.cells[1].text = f"{b_val['average']:.2f}"
        row.cells[2].text = b_val['level']
        row.cells[3].text = b_val['lawReference']

    doc.add_paragraph()

    # Capítulo IV - Recomendaciones y Plan de Acción
    h4 = doc.add_paragraph()
    r_h4 = h4.add_run("IV. RECOMENDACIONES LEGALES Y PLAN DE ACCIÓN REQUERIDO")
    r_h4.font.name = 'Arial'
    r_h4.font.size = Pt(14)
    r_h4.font.bold = True

    p_intro_actions = doc.add_paragraph()
    p_intro_actions.add_run("Con base en los ítems calificados como NO (0.00) o PARCIAL (0.50), "
                           "se han estructurado las siguientes tareas preventivas priorizadas de cumplimiento legal colombiano:")

    # Buscar recomendaciones
    # Core
    has_recs = False
    for block in COMPONENT_BLOCKS:
        for item in block['items']:
            for q in item['questions']:
                val = engine.answers.get(q['id'], 0.0)
                if val < 1.0:
                    has_recs = True
                    p_rec = doc.add_paragraph(style='List Bullet')
                    r_rec_bold = p_rec.add_run(f"[{block['name']}] / {item['name']}: ")
                    r_rec_bold.bold = True
                    p_rec.add_run(f"{q['recommendation']} (Pregunta: {q['text']})")

    # Transversales
    for block in TRANSVERSAL_BLOCKS:
        for q in block['questions']:
            val = engine.answers.get(q['id'], 0.0)
            if val < 1.0:
                has_recs = True
                p_rec = doc.add_paragraph(style='List Bullet')
                r_rec_bold = p_rec.add_run(f"[{block['name']}] (Base de Ley): ")
                r_rec_bold.bold = True
                p_rec.add_run(f"{q['recommendation']} (Pregunta: {q['text']})")

    if not has_recs:
        p_ok = doc.add_paragraph()
        p_ok.add_run("¡Excelente! No se registran vulnerabilidades activas o fallas operacionales en ninguno de los componentes evaluados.")

    # Firmas
    doc.add_paragraph("\n\n\n")
    table_firmas = doc.add_table(rows=1, cols=2)
    c_f = table_firmas.rows[0].cells
    c_f[0].text = "____________________________________\nResponsable de Calificación de Emergencias SG-SST\nLicencia de SST / Firma"
    c_f[1].text = "____________________________________\nRepresentante Legal / Gerente de Sede\nValidación del Presupuesto COE"

    try:
        doc.save(outfile)
        return True
    except Exception as e:
        print(f"Error al escribir el archivo Word {outfile}: {e}")
        return False

# Exportador a Excel usando openpyxl
def exportar_excel_xlsx(engine, outfile="Memoria_Tecnica_Vulnerabilidad.xlsx"):
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        print(f"\n[{Color.RED}AVISO{Color.RESET}] No se encuentra instalada la librería 'openpyxl'. "
              f"Instálela ejecutando: {Color.BOLD}pip install openpyxl{Color.RESET}")
        return False

    results = engine.calc_results()
    wb = openpyxl.Workbook()

    # 1. Hoja 1: Resumen Ejecutivo
    ws1 = wb.active
    ws1.title = "Resumen de Resultados"
    ws1.views.sheetView[0].showGridLines = True

    # Estilos
    font_title = Font(name="Arial", size=14, bold=True, color="FFFFFF")
    font_section = Font(name="Arial", size=11, bold=True, color="333333")
    font_bold = Font(name="Arial", size=10, bold=True)
    font_normal = Font(name="Arial", size=10)
    
    fill_slate = PatternFill(start_color="1E293B", end_color="1E293B", fill_type="solid")
    fill_teal = PatternFill(start_color="0D9488", end_color="0D9488", fill_type="solid")
    fill_header = PatternFill(start_color="F1F5F9", end_color="F1F5F9", fill_type="solid")

    thin_border = Border(
        left=Side(style='thin', color='CBD5E1'),
        right=Side(style='thin', color='CBD5E1'),
        top=Side(style='thin', color='CBD5E1'),
        bottom=Side(style='thin', color='CBD5E1')
    )

    # Banner Título
    ws1.merge_cells("A1:D1")
    ws1["A1"] = "INFORME EJECUTIVO DE VULNERABILIDAD"
    ws1["A1"].font = font_title
    ws1["A1"].fill = fill_slate
    ws1["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws1.row_dimensions[1].height = 40

    ws1["A3"] = "Sede:"
    ws1["A3"].font = font_bold
    ws1["B3"] = "Sede Principal Colombia / SG-SST"
    ws1["B3"].font = font_normal

    ws1["A4"] = "Evaluación:"
    ws1["A4"].font = font_bold
    ws1["B4"] = engine.threat_name
    ws1["B4"].font = font_normal

    ws1["A5"] = "Severidad:"
    ws1["A5"].font = font_bold
    ws1["B5"] = engine.threat_level
    ws1["B5"].font = font_normal

    # Tabla Core
    ws1["A7"] = "ANÁLISIS DE VULNERABILIDAD CORE (SEMÁFORO)"
    ws1["A7"].font = font_section
    
    headers_core = ["Componente Evaluado", "Calificación Promedio", "Vulnerabilidad", "Riesgo de Emergencia"]
    for col_idx, header in enumerate(headers_core, start=1):
        cell = ws1.cell(row=8, column=col_idx, value=header)
        cell.font = font_title
        cell.fill = fill_teal
        cell.alignment = Alignment(horizontal="center")

    row_pos = 9
    for b_id, b_val in results['core_results'].items():
        ws1.cell(row=row_pos, column=1, value=b_val['name']).font = font_bold
        ws1.cell(row=row_pos, column=2, value=b_val['average']).font = font_normal
        ws1.cell(row=row_pos, column=3, value=b_val['level']).font = font_normal
        
        status_label = "🔴 CRÍTICO" if b_val['color'] == 'red' else ("🟡 ADVERTENCIA" if b_val['color'] == 'yellow' else "🟢 PROTEGIDO")
        ws1.cell(row=row_pos, column=4, value=status_label).font = font_bold
        
        for c in range(1, 5):
            ws1.cell(row=row_pos, column=c).border = thin_border
        row_pos += 1

    # Tabla Resumen Diagnóstico
    row_pos += 2
    ws1.cell(row=row_pos, column=1, value="CONSOLIDADO FINAL").font = font_section
    row_pos += 1
    
    metrics = [
        ("Nivel de Vulnerabilidad Global", results['consolidated_vul_level']),
        ("Número de Rombos en Rojo (Vulnerabilidad)", results['red_vulnerability_rombos']),
        ("Rombos en Rojo Totales (Vulnerabilidad + Amenaza)", results['total_red_rombos']),
        ("Intersección de Riesgo Global", results['risk_level']),
        ("Descripción Operativa de Riesgo", results['risk_desc'])
    ]

    for label, val in metrics:
        cell_lbl = ws1.cell(row=row_pos, column=1, value=label)
        cell_lbl.font = font_bold
        cell_lbl.border = thin_border
        cell_v = ws1.cell(row=row_pos, column=2, value=val)
        cell_v.font = font_normal
        cell_v.border = thin_border
        ws1.merge_cells(start_row=row_pos, start_column=2, end_row=row_pos, end_column=4)
        row_pos += 1

    # Autosize columnas de la Hoja 1
    for col in ws1.columns:
        max_len = max(len(str(cell.value or '')) for cell in col)
        col_letter = get_column_letter(col[0].column)
        ws1.column_dimensions[col_letter].width = max(max_len + 3, 15)

    # 2. Hoja 2: Respuestas Detalladas (Core)
    ws2 = wb.create_sheet(title="Detalle Respuestas Core")
    ws2.views.sheetView[0].showGridLines = True

    headers_resp = ["Pregunta ID", "Componente", "Sub-bloque", "Pregunta de Evaluación", "Puntaje Registrado", "Acción Correctiva de Ley"]
    for col_idx, header in enumerate(headers_resp, start=1):
        cell = ws2.cell(row=1, column=col_idx, value=header)
        cell.font = font_title
        cell.fill = fill_slate
        cell.alignment = Alignment(horizontal="center")

    idx = 2
    for block in COMPONENT_BLOCKS:
        for item in block['items']:
            for q in item['questions']:
                val = engine.answers.get(q['id'], 0.0)
                ws2.cell(row=idx, column=1, value=q['id']).font = font_bold
                ws2.cell(row=idx, column=2, value=block['name']).font = font_normal
                ws2.cell(row=idx, column=3, value=item['name']).font = font_normal
                ws2.cell(row=idx, column=4, value=q['text']).font = font_normal
                
                cell_v = ws2.cell(row=idx, column=5, value=val)
                cell_v.font = font_bold
                cell_v.alignment = Alignment(horizontal="center")
                if val == 1.0:
                    cell_v.fill = PatternFill(start_color="DCFCE7", end_color="DCFCE7", fill_type="solid") # soft green
                elif val == 0.5:
                    cell_v.fill = PatternFill(start_color="FEF3C7", end_color="FEF3C7", fill_type="solid") # soft yellow
                else:
                    cell_v.fill = PatternFill(start_color="FEE2E2", end_color="FEE2E2", fill_type="solid") # soft red

                ws2.cell(row=idx, column=6, value=q['recommendation']).font = font_normal
                
                for c in range(1, 7):
                    ws2.cell(row=idx, column=c).border = thin_border
                idx += 1

    for col in ws2.columns:
        col_letter = get_column_letter(col[0].column)
        ws2.column_dimensions[col_letter].width = 18
    ws2.column_dimensions['D'].width = 50
    ws2.column_dimensions['F'].width = 50

    # 3. Hoja 3: Respuestas Detalladas (Transversales)
    ws3 = wb.create_sheet(title="Transversales Ley")
    ws3.views.sheetView[0].showGridLines = True

    for col_idx, header in enumerate(headers_resp, start=1):
        cell = ws3.cell(row=1, column=col_idx, value=header)
        cell.font = font_title
        cell.fill = fill_slate
        cell.alignment = Alignment(horizontal="center")

    idx = 2
    for block in TRANSVERSAL_BLOCKS:
        for q in block['questions']:
            val = engine.answers.get(q['id'], 0.0)
            ws3.cell(row=idx, column=1, value=q['id']).font = font_bold
            ws3.cell(row=idx, column=2, value=block['name']).font = font_normal
            ws3.cell(row=idx, column=3, value=block['lawReference']).font = font_normal
            ws3.cell(row=idx, column=4, value=q['text']).font = font_normal
            
            cell_v = ws3.cell(row=idx, column=5, value=val)
            cell_v.font = font_bold
            cell_v.alignment = Alignment(horizontal="center")
            if val == 1.0:
                cell_v.fill = PatternFill(start_color="DCFCE7", end_color="DCFCE7", fill_type="solid")
            elif val == 0.5:
                cell_v.fill = PatternFill(start_color="FEF3C7", end_color="FEF3C7", fill_type="solid")
            else:
                cell_v.fill = PatternFill(start_color="FEE2E2", end_color="FEE2E2", fill_type="solid")

            ws3.cell(row=idx, column=6, value=q['recommendation']).font = font_normal
            
            for c in range(1, 7):
                ws3.cell(row=idx, column=c).border = thin_border
            idx += 1

    for col in ws3.columns:
        col_letter = get_column_letter(col[0].column)
        ws3.column_dimensions[col_letter].width = 18
    ws3.column_dimensions['D'].width = 50
    ws3.column_dimensions['F'].width = 50

    try:
        wb.save(outfile)
        return True
    except Exception as e:
        print(f"Error al escribir el archivo Excel {outfile}: {e}")
        return False

# Wizard interactivo del terminal
def main():
    # Inicializar motor de evaluación
    engine = EvaluacionEngine()
    
    # Cargar sesión previa si existe
    if engine.load_session():
        print(f"{Color.GREEN}✔ Se cargó automáticamente una sesión de trabajo anterior.{Color.RESET}")
    else:
        # Pestaña predeterminada
        engine.load_preset(1) # PyME por defecto

    while True:
        os.system('cls' if os.name == 'nt' else 'clear')
        print(f"{Color.CYAN}{Color.BOLD}================================================================")
        print("          CALCULADOR DE VULNERABILIDAD POR ROMBOS SG-SST        ")
        print(f"================================================================{Color.RESET}")
        
        # Estado actual de cálculo
        res = engine.calc_results()
        
        print(f"{Color.BOLD}Escenario de Amenaza:{Color.RESET} {Color.WHITE}{engine.threat_name}{Color.RESET}")
        print(f"{Color.BOLD}Nivel de Probabilidad:{Color.RESET} {engine.threat_level}")
        print(f"{Color.BOLD}Descripción:{Color.RESET} {Color.GRAY}{engine.threat_desc[:65]}...{Color.RESET}\n")

        print(f"{Color.BOLD}--- Estado de Calificación Core ---{Color.RESET}")
        for b_id, b_val in res['core_results'].items():
            print(f" ⚫ {b_val['name']}: {b_val['color_ansi']}{b_val['level']}{Color.RESET} (Promedio: {b_val['average']})")

        print(f"\n{Color.BOLD}--- Estado de Bloques Transversales de Ley ---{Color.RESET}")
        for b_id, b_val in res['transversal_results'].items():
            print(f" ✦ {b_val['name']}: {b_val['color_ansi']}{b_val['level']}{Color.RESET} (Promedio: {b_val['average']})")

        # Intersección final
        print(f"\n{Color.BOLD}================================================================")
        print(f" RIESGO INTERCEPTADO: {res['risk_color_ansi']}{Color.BOLD}RIESGO {res['risk_level']}{Color.RESET}")
        print(f" {res['risk_color_ansi']}{res['risk_desc']}{Color.RESET}")
        print(f" Rombos de Vulnerabilidad en Rojo: {Color.RED}{res['red_vulnerability_rombos']}/3{Color.RESET} | Rombos en Rojo Totales con la Amenaza: {Color.RED}{res['total_red_rombos']}/4{Color.RESET}")
        print(f"================================================================{Color.RESET}\n")

        # Menú
        print(f"{Color.BOLD}Seleccione una opción:{Color.RESET}")
        print(" [1] Configurar Amenaza de Riesgo (Nombre, descripción, probabilidad)")
        print(" [2] Cargar Plantilla Predefinida (Industrial 🟢, PyME 🟡, Edificio Crítico 🔴)")
        print(" [3] Evaluar Componentes Core (Personas, Recursos, Sistemas)")
        print(" [4] Evaluar Módulos Transversales de Ley (PESV, Psicosocial, Alturas, etc.)")
        print(" [5] Descargar Memoria Técnica en Excel (.xlsx)")
        print(" [6] Generar Plan de Emergencias en Word (.docx)")
        print(" [7] Dibujar e Insertar Imagen del Diamante (.png)")
        print(" [8] Limpiar respuestas a cero (Comenzar de nuevo)")
        print(" [9] Guardar progreso manual")
        print(" [0] Salir de la aplicación")
        
        opt = input(f"\n{Color.BOLD}Opción > {Color.RESET}").strip()

        if opt == '1':
            os.system('cls' if os.name == 'nt' else 'clear')
            print(f"{Color.CYAN}{Color.BOLD}--- CARACTERIZACIÓN DE LA AMENAZA ---{Color.RESET}\n")
            print("Escenarios sugeridos:")
            for idx, th in enumerate(STANDARD_THREATS, start=1):
                print(f" {idx}. {th['name']} (Nivel Típico: {th['level']})")
            
            print("\nEscriba el número de la amenaza recomendada para auto-cargar sus detalles, o presione ENTER para personalizar:")
            th_sel = input("Sugerencia > ").strip()
            
            if th_sel.isdigit() and 1 <= int(th_sel) <= len(STANDARD_THREATS):
                th = STANDARD_THREATS[int(th_sel) - 1]
                engine.threat_name = th['name']
                engine.threat_level = th['level']
                engine.threat_desc = th['description']
            else:
                engine.threat_name = input(f"{Color.BOLD}Nombre de la Amenaza: {Color.RESET}").strip() or engine.threat_name
                engine.threat_desc = input(f"{Color.BOLD}Descripción del escenario: {Color.RESET}").strip() or engine.threat_desc
                
            print(f"\nNivel de probabilidad de la amenaza:\n 1 = {Color.GREEN}POSIBLE (Verde){Color.RESET}\n 2 = {Color.YELLOW}PROBABLE (Amarillo){Color.RESET}\n 3 = {Color.RED}INMINENTE (Rojo){Color.RESET}")
            p_sel = input("Probabilidad [1/2/3] > ").strip()
            if p_sel == '1':
                engine.threat_level = "POSIBLE"
            elif p_sel == '2':
                engine.threat_level = "PROBABLE"
            elif p_sel == '3':
                engine.threat_level = "INMINENTE"

            engine.save_session()

        elif opt == '2':
            os.system('cls' if os.name == 'nt' else 'clear')
            print(f"{Color.CYAN}{Color.BOLD}--- CARGAR PRESET ---{Color.RESET}\n")
            for idx, pr in enumerate(EVALUATION_PRESETS, start=1):
                print(f"  {idx}. {pr['name']}")
                print(f"     {Color.GRAY}{pr['threat_name']} / {pr['threat_level']}{Color.RESET}")
            
            sel = input(f"\nSeleccione [1-3] > ").strip()
            if sel.isdigit() and 1 <= int(sel) <= 3:
                engine.load_preset(int(sel) - 1)
                engine.save_session()
                print(f"\n{Color.GREEN}✔ Preset cargado con éxito. Respuestas autocompletadas.{Color.RESET}")
                input("Presione ENTER para continuar...")

        elif opt == '3':
            # Evaluador Core
            for block in COMPONENT_BLOCKS:
                os.system('cls' if os.name == 'nt' else 'clear')
                print(f"{Color.CYAN}{Color.BOLD}--- {block['name'].upper()} ---{Color.RESET}")
                print("Escriba para cada pregunta: \n 1 = SÍ (1.0) \n 2 = PARCIAL (0.5) \n 0 = NO (0.0)\n Pulse ENTER para mantener respuesta actual o presione 'S' para salir.\n")
                
                for item in block['items']:
                    print(f"\n{Color.BOLD}» Subtema: {item['name']}{Color.RESET}")
                    print(f"  {Color.GRAY}{item['description']}{Color.RESET}")
                    for q in item['questions']:
                        curr = engine.answers.get(q['id'], None)
                        curr_label = f"SÍ (1.0)" if curr == 1.0 else (f"PARCIAL (0.5)" if curr == 0.5 else (f"NO (0.0)" if curr == 0.0 else "Sin calificar"))
                        
                        print(f"\n✦ {q['text']}")
                        print(f"  {Color.CYAN}Valor actual: {curr_label}{Color.RESET}")
                        ans = input(f"  Respuesta [1 / 2 / 0 / S] > ").strip().upper()
                        
                        if ans == 'S':
                            break
                        elif ans == '1':
                            engine.set_answer(q['id'], 1.0)
                        elif ans == '2':
                            engine.set_answer(q['id'], 0.5)
                        elif ans == '0':
                            engine.set_answer(q['id'], 0.0)
                    if ans == 'S':
                        break
                engine.save_session()

        elif opt == '4':
            # Evaluador Transversal
            for block in TRANSVERSAL_BLOCKS:
                os.system('cls' if os.name == 'nt' else 'clear')
                print(f"{Color.CYAN}{Color.BOLD}--- EVALUACIÓN TRANSVERSAL DE LEY ---{Color.RESET}")
                print(f"{Color.BOLD}Bloque: {block['name']}{Color.RESET}")
                print(f"Norma: {Color.YELLOW}{block['lawReference']}{Color.RESET}")
                print("Escriba para cada pregunta: \n 1 = SÍ (1.0) \n 2 = PARCIAL (0.5) \n 0 = NO (0.0)\n Pulse ENTER para mantener o 'S' para salir.\n")
                
                for q in block['questions']:
                    curr = engine.answers.get(q['id'], None)
                    curr_label = "SÍ (1.0)" if curr == 1.0 else ("PARCIAL (0.5)" if curr == 0.5 else ("NO (0.0)" if curr == 0.0 else "Sin calificar"))
                    
                    print(f"\n✦ {q['text']}")
                    print(f"  {Color.CYAN}Valor actual: {curr_label}{Color.RESET}")
                    ans = input(f"  Respuesta [1 / 2 / 0 / S] > ").strip().upper()
                    
                    if ans == 'S':
                        break
                    elif ans == '1':
                        engine.set_answer(q['id'], 1.0)
                    elif ans == '2':
                        engine.set_answer(q['id'], 0.5)
                    elif ans == '0':
                        engine.set_answer(q['id'], 0.0)
                if ans == 'S':
                    break
                engine.save_session()

        elif opt == '5':
            print("\nGenerando memoria analítica en Excel...")
            f_n = "Memoria_Tecnica_Vulnerabilidad.xlsx"
            if exportar_excel_xlsx(engine, f_n):
                print(f"{Color.GREEN}✔ Excel exportado exitosamente como '{f_n}'. Puede abrirlo en Microsoft Excel.{Color.RESET}")
            input("\nPresione ENTER para retornar...")

        elif opt == '6':
            print("\nGenerando plan reglamentario de emergencias en Word...")
            f_n = "Plan_Prevencion_Emergencias.docx"
            if exportar_word_docx(engine, f_n):
                print(f"{Color.GREEN}✔ Plan de Emergencia en Word redactado con éxito como '{f_n}'. Contiene introducción, tablas analíticas formalizadas y el gráfico del diamante.{Color.RESET}")
            input("\nPresione ENTER para retornar...")

        elif opt == '7':
            print("\nGenerando imagen vectorial del semáforo por rombos...")
            f_n = "diamante_riesgo.png"
            if generar_imagen_diamante_pil(engine, f_n):
                print(f"{Color.GREEN}✔ Imagen renderizada de alta resolución exportada como '{f_n}' con colores dinámicos.{Color.RESET}")
            input("\nPresione ENTER para retornar...")

        elif opt == '8':
            confirm = input(f"\n¿Está seguro de que desea limpiar todas las respuestas a cero? [S/N] > ").strip().upper()
            if confirm == 'S':
                engine.answers = {}
                engine.save_session()
                print(f"{Color.GREEN}✔ Respuestas vaciadas con éxito.{Color.RESET}")
            input("Presione ENTER para continuar...")

        elif opt == '9':
            if engine.save_session():
                print(f"\n{Color.GREEN}✔ Progreso guardado de manera manual en '{SESSION_FILE}'. El archivo se auto-cargará en la próxima ejecución.{Color.RESET}")
            input("Presione ENTER para continuar...")

        elif opt == '0':
            engine.save_session()
            print(f"\n{Color.CYAN}¡Gracias por utilizar el Calculador de Rombos SG-SST! Sesión guardada de forma segura.{Color.RESET}")
            break

if __name__ == '__main__':
    main()
